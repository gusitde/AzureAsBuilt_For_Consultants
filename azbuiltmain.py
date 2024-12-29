import os
import logging
from azure.identity import DefaultAzureCredential
from azure.mgmt.resource import ResourceManagementClient
from azure.mgmt.network import NetworkManagementClient
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure logging
log_filename = 'asbuiltlogs.txt'
logging.basicConfig(level=logging.DEBUG,  # Set to DEBUG to capture all levels of logs
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    handlers=[
                        logging.FileHandler(log_filename, mode='w'),  # Write mode to overwrite each run
                        logging.StreamHandler()
                    ])

# Mapping of resource types to Azure service names and descriptions
RESOURCE_TYPE_DETAILS = {
    "Microsoft.CognitiveServices/accounts": ("Azure Cognitive Services", "Provides AI and machine learning services."),
    "Microsoft.Compute/virtualMachines": ("Azure Virtual Machines", "Scalable computing resources for running applications."),
    "Microsoft.Network/virtualNetworks": ("Azure Virtual Networks", "Enables secure connections between Azure services."),
    "Microsoft.Storage/storageAccounts": ("Azure Storage Accounts", "Scalable cloud storage solutions."),
    "Microsoft.Web/sites": ("Azure App Service", "Platform for building and hosting web apps."),
    "Microsoft.Sql/servers": ("Azure SQL Database", "Managed database service for SQL Server."),
    "Microsoft.KeyVault/vaults": ("Azure Key Vault", "Securely stores and manages access to secrets."),
    "Microsoft.ContainerRegistry/registries": ("Azure Container Registry", "Stores and manages container images."),
    "Microsoft.Kubernetes/connectedClusters": ("Azure Arc-enabled Kubernetes", "Manages Kubernetes clusters across environments."),
    "Microsoft.Network/publicIPAddresses": ("Azure Public IP Addresses", "Provides static and dynamic public IP addresses."),
    "Microsoft.Network/networkSecurityGroups": ("Azure Network Security Groups", "Controls inbound and outbound traffic to resources."),
    "Microsoft.Network/loadBalancers": ("Azure Load Balancers", "Distributes traffic among multiple servers."),
    "Microsoft.Network/applicationGateways": ("Azure Application Gateways", "Manages application delivery and load balancing."),
    "Microsoft.Network/dnszones": ("Azure DNS Zones", "Hosts DNS domains and manages DNS records."),
    "Microsoft.Network/expressRouteCircuits": ("Azure ExpressRoute Circuits", "Private connections between on-premises networks and Azure."),
    "Microsoft.Network/virtualNetworkGateways": ("Azure Virtual Network Gateways", "Connects on-premises networks to Azure VNets."),
    "Microsoft.Network/routeTables": ("Azure Route Tables", "Defines routes for network traffic."),
    "Microsoft.ContainerInstance/containerGroups": ("Azure Container Instances", "Runs containers without managing servers."),
    "Microsoft.ContainerService/managedClusters": ("Azure Kubernetes Service (AKS)", "Managed Kubernetes service for containerized applications."),
    "Microsoft.DocumentDB/databaseAccounts": ("Azure Cosmos DB", "Globally distributed multi-model database service."),
    "Microsoft.EventHub/namespaces": ("Azure Event Hubs", "Big data streaming platform and event ingestion service."),
    "Microsoft.Insights/components": ("Azure Application Insights", "Monitors and diagnoses application performance issues."),
    "Microsoft.Logic/workflows": ("Azure Logic Apps", "Automates workflows and integrates apps and data."),
    "Microsoft.MachineLearningServices/workspaces": ("Azure Machine Learning", "Platform for building and deploying machine learning models."),
    "Microsoft.ManagedIdentity/userAssignedIdentities": ("Azure Managed Identities", "Provides identity management for Azure resources."),
    "Microsoft.OperationalInsights/workspaces": ("Azure Log Analytics", "Collects and analyzes log data."),
    "Microsoft.Relay/namespaces": ("Azure Relay", "Enables hybrid applications by bridging on-premises and cloud environments."),
    "Microsoft.Search/searchServices": ("Azure Cognitive Search", "Search-as-a-service for building search experiences."),
    "Microsoft.ServiceBus/namespaces": ("Azure Service Bus", "Fully managed enterprise message broker."),
    "Microsoft.SignalRService/signalr": ("Azure SignalR Service", "Real-time messaging service for web applications."),
    "Microsoft.Sql/servers/databases": ("Azure SQL Databases", "Managed relational database service."),
    "Microsoft.StreamAnalytics/streamingjobs": ("Azure Stream Analytics", "Real-time data processing service."),
    "Microsoft.Synapse/workspaces": ("Azure Synapse Analytics", "Analytics service that brings together big data and data warehousing."),
    "Microsoft.Web/serverfarms": ("Azure App Service Plans", "Plans for hosting web apps, mobile apps, and API apps."),
    # Add any other mappings as needed
}

# Define specific headers for each Azure service type
SERVICE_HEADERS = {
    "Azure Virtual Machines": ["Name", "Resource Group", "Location", "Size", "OS Type", "Tags"],
    "Azure Virtual Networks": ["Name", "Resource Group", "Location", "Address Space", "Tags"],
    "Azure Storage Accounts": ["Name", "Resource Group", "Location", "SKU", "Access Tier", "Tags"],
    "Azure App Service": ["Name", "Resource Group", "Location", "App Service Plan", "State", "Tags"],
    "Azure SQL Database": ["Name", "Resource Group", "Location", "Database Edition", "Service Objective", "Tags"],
    # Add more custom headers for other service types as needed
}

def load_azure_data(subscription_ids):
    """Load data from Azure for given subscription IDs."""
    logging.info("Loading Azure data for subscription IDs.")
    credential = DefaultAzureCredential()
    resource_clients = [{"client": ResourceManagementClient(credential, sub_id), "subscription_id": sub_id} for sub_id in subscription_ids]
    network_clients = [{"client": NetworkManagementClient(credential, sub_id), "subscription_id": sub_id} for sub_id in subscription_ids]
    return resource_clients, network_clients

def fetch_resources(resource_clients):
    """Fetch all resources for given clients and organize them by type."""
    all_resources = {}
    for client_info in resource_clients:
        client = client_info["client"]
        sub_id = client_info["subscription_id"]
        try:
            logging.info(f"Fetching resources for client with subscription ID {sub_id}.")
            resources = client.resources.list()
            for resource in resources:
                resource_type = resource.type
                if resource_type not in all_resources:
                    all_resources[resource_type] = []
                all_resources[resource_type].append(resource.as_dict())
        except Exception as e:
            logging.error(f"Error fetching data for client with subscription ID {sub_id}: {e}")
    return all_resources

def fetch_network_details(network_clients):
    """Fetch details for specific network resources like VNets."""
    network_details = {}
    for client_info in network_clients:
        client = client_info["client"]
        sub_id = client_info["subscription_id"]
        try:
            logging.info(f"Fetching network details for client with subscription ID {sub_id}.")
            vnets = client.virtual_networks.list_all()
            network_details['virtualNetworks'] = [vnet.as_dict() for vnet in vnets]
        except Exception as e:
            logging.error(f"Error fetching network details for client with subscription ID {sub_id}: {e}")
    return network_details

def process_resource_data(resources, network_details):
    """Process resource data to extract relevant information."""
    sections = []
    counts = {
        "subscriptions": len(resources),
        "resource_groups": set(),
        "virtual_machines": 0,
        "disks": 0,
        "storage_accounts": 0,
        "vnets": 0,
    }

    for resource_type, resource_list in resources.items():
        logging.debug(f"Processing resource type: {resource_type}")
        service_name, service_description = RESOURCE_TYPE_DETAILS.get(resource_type, (resource_type, "Description not available."))
        logging.debug(f"Service name: {service_name}, Description: {service_description}")
        section_content = []
        for resource in resource_list:
            # Update counts
            counts["resource_groups"].add(resource.get('resourceGroup', 'N/A'))
            if resource_type == "Microsoft.Compute/virtualMachines":
                counts["virtual_machines"] += 1
            elif resource_type == "Microsoft.Compute/disks":
                counts["disks"] += 1
            elif resource_type == "Microsoft.Storage/storageAccounts":
                counts["storage_accounts"] += 1
            elif resource_type == "Microsoft.Network/virtualNetworks":
                counts["vnets"] += 1

            # Fetch additional details for each resource
            resource_details = {
                "Name": resource.get('name', 'N/A'),
                "Resource Group": resource.get('resourceGroup', 'N/A'),
                "Location": resource.get('location', 'N/A'),
                "Kind": resource.get('kind', 'N/A'),
                "SKU": resource.get('sku', {}).get('name', 'N/A'),
                "Tags": resource.get('tags', 'N/A'),
                "ID": resource.get('id', 'N/A')
            }

            # Add address space for VNets
            if resource_type == "Microsoft.Network/virtualNetworks":
                resource_details["Address Space"] = ', '.join(resource.get('addressSpace', {}).get('addressPrefixes', []))

            section_content.append(resource_details)
        sections.append({
            "title": f"Service: {service_name}",
            "description": service_description,
            "content": section_content
        })

    # Convert the set of resource groups to a count
    counts["resource_groups"] = len(counts["resource_groups"])

    logging.info(f"Processed resource data: {counts}")
    return sections, counts

def remove_empty_columns(headers, content):
    """Remove columns that are entirely empty from headers and content."""
    non_empty_headers = []
    non_empty_content = []
    for idx, header in enumerate(headers):
        if any(item.get(header, 'N/A') != 'N/A' for item in content):
            non_empty_headers.append(header)

    for item in content:
        non_empty_item = {k: v for k, v in item.items() if k in non_empty_headers}
        non_empty_content.append(non_empty_item)

    return non_empty_headers, non_empty_content

def add_custom_styles(doc):
    """Add custom styles to the Word document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Aptos'  # Ensure 'Aptos' font is installed
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)  # Black color

def set_table_borders(table):
    """Set borders for all cells in a table and add outside borders."""
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)

def format_table_header(table, headers):
    """Format table header with bold letters, white font color, and a light blue background."""
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        # Set the text format
        cell_para = hdr_cells[idx].paragraphs[0]
        cell_run = cell_para.runs[0]
        cell_run.bold = True
        cell_run.font.color.rgb = RGBColor(255, 255, 255)  # White font
        cell_run.font.size = Pt(12)
        
        # Set the cell background color
        cell_tcPr = hdr_cells[idx]._element.get_or_add_tcPr()
        cell_shading = OxmlElement('w:shd')
        cell_shading.set(qn('w:fill'), '87CEEB')  # Light sky blue background
        cell_tcPr.append(cell_shading)

def generate_document(sections, counts, filename='asbuilt.docx'):
    """Generate a Word document with the given sections and counts."""
    doc = Document()
    add_custom_styles(doc)

    # Title Page
    title_page = doc.add_paragraph()
    title_page.add_run("As-Built Document").bold = True
    title_page.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\n\n\n\n\n")  # Add some space

    # Summary Section
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(
        "This As-Built Document provides a comprehensive overview of the current state of Azure resources "
        "within the specified subscription IDs. It includes detailed information about various services, "
        "such as Virtual Machines, Storage Accounts, Virtual Networks, and more. Each section contains "
        "a description of the service, a table of key resource attributes, and unique resource IDs."
    )

    # Total Counts Section
    doc.add_heading('Total Counts', level=1)
    doc.add_paragraph(f"Subscriptions: {counts['subscriptions']}")
    doc.add_paragraph(f"Resource Groups: {counts['resource_groups']}")
    doc.add_paragraph(f"Virtual Machines: {counts['virtual_machines']}")
    doc.add_paragraph(f"Disks: {counts['disks']}")
    doc.add_paragraph(f"Storage Accounts: {counts['storage_accounts']}")
    doc.add_paragraph(f"Virtual Networks: {counts['vnets']}")

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc = doc.add_paragraph()
    toc_run = toc.add_run()
    for i, section in enumerate(sections):
        toc_run.add_text(f"{i + 1}. {section['title']} ................... ")
        toc_run.add_text(f"Page {i + 6}")  # Assuming the title page, summary, counts, and TOC take up 5 pages
        toc_run.add_break()

    doc.add_page_break()

    # Content Sections
    for section in sections:
        doc.add_heading(section['title'], level=1)
        doc.add_paragraph(section['description'])

        # Determine headers based on service type
        service_name = section['title'].replace("Service: ", "")
        headers = SERVICE_HEADERS.get(service_name, ["Name", "Resource Group", "Location", "Kind", "SKU", "Tags"])

        # Add address space for VNets
        if service_name == "Azure Virtual Networks":
            headers = ["Name", "Resource Group", "Location", "Address Space", "Tags"]

        # Remove empty columns
        headers, content = remove_empty_columns(headers, section['content'])

        # Create a table for each section with customized headers
        table = doc.add_table(rows=1, cols=len(headers))
        format_table_header(table, headers)

        for item in content:
            row_cells = table.add_row().cells
            for idx, header in enumerate(headers):
                row_cells[idx].text = item.get(header, 'N/A')

        set_table_borders(table)

        # Add resource IDs under the table
        for item in section['content']:
            resource_id = item.get('ID', 'N/A')
            logging.info(f"Resource ID: {resource_id}")  # Log the ID for debugging
            id_paragraph = doc.add_paragraph()
            id_run = id_paragraph.add_run(f"ID: {resource_id}")
            id_run.bold = True

        doc.add_paragraph("\n")  # Add a space between sections

    doc.save(filename)
    logging.info(f"Document saved as {filename}")

def main():
    logging.info("Starting the As-Built Document generation process.")
    subscription_ids = os.getenv('AZURE_SUBSCRIPTION_IDS', '5514d116-97eb-4cfc-927f-b03826fcc9cc').split(',')
    resource_clients, network_clients = load_azure_data(subscription_ids)
    all_resources = fetch_resources(resource_clients)
    network_details = fetch_network_details(network_clients)
    sections, counts = process_resource_data(all_resources, network_details)
    generate_document(sections, counts)
    logging.info("As-Built Document generation process completed.")

if __name__ == "__main__":
    main()
